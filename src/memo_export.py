"""Convert memo markdown to DOCX."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def extract_memo_from_text(text: str) -> str:
    """Pull memo body from Claude response markers."""
    markers = [
        ("--- MEMO START ---", "--- MEMO END ---"),
        ("```memo", "```"),
        ("```markdown", "```"),
    ]
    for start, end in markers:
        if start in text:
            block = text.split(start, 1)[1]
            if end in block:
                return block.split(end, 1)[0].strip()
            return block.strip()

    # Fallback: if response looks like a full memo (has headings), use as-is
    if re.search(r"^#\s+", text, re.MULTILINE) or "Executive Summary" in text:
        return text.strip()
    return ""


def markdown_to_docx(markdown_text: str, output_path: Path, *, title: str = "") -> Path:
    output_path = Path(output_path)
    if output_path.exists():
        raise FileExistsError(f"Refusing to overwrite: {output_path}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    lines = markdown_text.splitlines()
    i = 0
    doc_title = title
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            heading = stripped[2:].strip()
            if not doc_title:
                doc_title = heading
                h = doc.add_heading(heading, 0)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_heading(heading, level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith(("- ", "* ")):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", stripped), style="List Number")
        else:
            para_lines = [stripped]
            i += 1
            while i < len(lines) and lines[i].strip() and not _is_block_start(lines[i]):
                para_lines.append(lines[i].strip())
                i += 1
            p = doc.add_paragraph(" ".join(para_lines))
            p.paragraph_format.space_after = Pt(8)
            i -= 1

        i += 1

    if not doc_title:
        doc.add_heading("BCA Technical Memorandum", 0)

    doc.save(str(output_path))
    return output_path


def _is_block_start(line: str) -> bool:
    s = line.strip()
    return bool(
        s.startswith("#")
        or s.startswith("- ")
        or s.startswith("* ")
        or re.match(r"^\d+\.\s", s)
    )
