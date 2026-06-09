"""Extract text from uploaded project documents."""

from __future__ import annotations

import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md", ".csv", ".xlsx", ".xlsm"}
MAX_CHARS_PER_FILE = 80_000
MAX_CHARS_TOTAL = 200_000


class DocumentIngestError(Exception):
    """Raised when a file cannot be processed."""


def extract_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise DocumentIngestError(
            f"Unsupported file type: {suffix}. Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    if suffix == ".pdf":
        return _extract_pdf(data)
    if suffix in (".docx", ".doc"):
        return _extract_docx(data)
    if suffix in (".txt", ".md", ".csv"):
        return _decode_text(data)
    if suffix in (".xlsx", ".xlsm"):
        return _extract_xlsx(data)

    raise DocumentIngestError(f"Unhandled extension: {suffix}")


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentIngestError("Could not decode text file")


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentIngestError("PDF support requires pypdf") from exc

    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)
    if not parts:
        raise DocumentIngestError("PDF contained no extractable text")
    return "\n\n".join(parts)


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentIngestError("DOCX support requires python-docx") from exc

    doc = Document(io.BytesIO(data))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    if not parts:
        raise DocumentIngestError("DOCX contained no extractable text")
    return "\n\n".join(parts)


def _extract_xlsx(data: bytes) -> str:
    return extract_workbook_structure(data, max_rows_per_sheet=120)


def extract_workbook_structure(
    data: bytes,
    *,
    max_rows_per_sheet: int | None = 120,
    include_formulas: bool = True,
) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise DocumentIngestError("Excel support requires openpyxl") from exc

    wb = load_workbook(
        io.BytesIO(data),
        read_only=True,
        data_only=not include_formulas,
    )
    parts: list[str] = []
    for sheet in wb.worksheets:
        parts.append(f"### Sheet: {sheet.title}")
        row_count = 0
        for row in sheet.iter_rows(values_only=True):
            if max_rows_per_sheet is not None and row_count >= max_rows_per_sheet:
                parts.append(f"[... {sheet.title}: additional rows omitted ...]")
                break
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append("\t".join(cells))
                row_count += 1
    wb.close()
    if len(parts) <= 1:
        raise DocumentIngestError("Spreadsheet contained no extractable data")
    return "\n\n".join(parts)


def _truncate(text: str, limit: int) -> tuple[str, bool]:
    if len(text) <= limit:
        return text, False
    return text[:limit] + "\n\n[... truncated for context limit ...]", True


def ingest_uploaded_files(
    files: list[tuple[str, bytes]],
    *,
    max_chars_per_file: int = MAX_CHARS_PER_FILE,
    max_chars_total: int = MAX_CHARS_TOTAL,
) -> tuple[str, list[dict[str, str]], list[str]]:
    if not files:
        return "", [], []

    records: list[dict[str, str]] = []
    warnings: list[str] = []
    combined_parts: list[str] = []
    total_chars = 0

    for filename, data in files:
        try:
            raw = extract_text(filename, data)
        except DocumentIngestError as exc:
            warnings.append(f"{filename}: {exc}")
            continue

        text, file_truncated = _truncate(raw, max_chars_per_file)
        if file_truncated:
            warnings.append(f"{filename}: truncated to {max_chars_per_file:,} characters")

        header = f"--- FILE: {filename} ({len(text)} chars) ---"
        combined_parts.append(f"{header}\n{text}")
        records.append({"filename": filename, "text": text, "char_count": str(len(text))})
        total_chars += len(text)

    combined = "\n\n".join(combined_parts)
    combined, total_truncated = _truncate(combined, max_chars_total)
    if total_truncated:
        warnings.append(f"Combined documents truncated to {max_chars_total:,} characters")

    return combined, records, warnings
